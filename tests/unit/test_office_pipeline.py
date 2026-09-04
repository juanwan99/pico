"""T-OFFICE-KERNEL: spec → render → inspect → addressable edit → fail-closed."""

from __future__ import annotations

import io
import sys
from pathlib import Path

import pytest

ROOT = Path(__file__).resolve().parents[2]
sys.path.insert(0, str(ROOT / "services" / "orchestrator"))

from pico_orchestrator.artifact_types import is_valid_ooxml_package
from pico_orchestrator.document_generators import build_docx_document, build_pptx_document
from pico_orchestrator.office.edit import edit_docx_bytes
from pico_orchestrator.office.inspect import inspect_office_bytes
from pico_orchestrator.office.qa import verify_office_bytes
from pico_orchestrator.office.render import render_spec
from pico_orchestrator.office.spec import parse_spec
from pico_orchestrator.true_pi.config import ALLOWED_GATEWAY_TOOLS

ONE_PNG = (
    b"\x89PNG\r\n\x1a\n"
    b"\x00\x00\x00\rIHDR\x00\x00\x00\x01\x00\x00\x00\x01\x08\x02\x00\x00\x00\x90wS\xde"
    b"\x00\x00\x00\x0cIDATx\x9cc``\x00\x00\x00\x04\x00\x01\xa3\x0a\x0d\xe4"
    b"\x00\x00\x00\x00IEND\xaeB`\x82"
)


def test_spec_docx_table_roundtrip():
    spec = parse_spec(
        {
            "schema": "pico.office.spec/v1",
            "kind": "docx",
            "title": "通知",
            "marker": "m-table",
            "blocks": [
                {"type": "heading", "text": "本周安排", "level": 0},
                {"type": "para", "text": "下面是分组表。"},
                {
                    "type": "table",
                    "rows": [["组", "人"], ["甲", "李老师"], ["乙", "王老师"]],
                },
            ],
        }
    )
    raw = render_spec(spec)
    assert is_valid_ooxml_package(raw, ".docx")
    outline = inspect_office_bytes(raw, ".docx")
    assert outline["tables"] == 1
    table = next(u for u in outline["units"] if u["kind"] == "table")
    assert table["rows"] == 3
    assert "组" in table["headers"]
    assert any("甲" in [str(c) for c in row] for row in table["preview"])


def test_spec_pptx_image_roundtrip():
    spec = parse_spec(
        {
            "kind": "pptx",
            "title": "汇报",
            "blocks": [
                {"type": "slide", "title": "封面", "bullets": ["要点一"]},
                {
                    "type": "slide",
                    "title": "配图",
                    "bullets": ["见图"],
                    "image_artifact_id": "img-1",
                },
                {"type": "slide", "title": "结尾", "bullets": ["谢谢"]},
            ],
        }
    )
    raw = render_spec(spec, images={"img-1": ONE_PNG})
    assert is_valid_ooxml_package(raw, ".pptx")
    outline = inspect_office_bytes(raw, ".pptx")
    assert outline["slides"] == 3
    assert outline["images"] >= 1
    pictured = next(u for u in outline["units"] if u["index"] == 2)
    assert pictured["images"] >= 1


def test_plain_generate_still_ooxml():
    body = ("这是完整通知正文。" * 20) + "\n\n第二段继续写清楚事项。"
    raw = build_docx_document(title="通知.docx", marker="mk-plain", body=body)
    assert is_valid_ooxml_package(raw, ".docx")
    outline = inspect_office_bytes(raw, ".docx")
    assert outline["paragraphs"] >= 2


def test_path_b_edit_keeps_other_paragraphs():
    from docx import Document

    doc = Document()
    doc.add_paragraph("第一段保持")
    doc.add_paragraph("第二段要改")
    doc.add_paragraph("第三段也在")
    buf = io.BytesIO()
    doc.save(buf)
    raw = buf.getvalue()
    edited = edit_docx_bytes(raw, paragraph_index=2, text="第二段已改")
    outline = inspect_office_bytes(edited, ".docx")
    texts = [u["text"] for u in outline["units"] if u["kind"] == "para"]
    assert "第一段保持" in texts
    assert "第二段已改" in texts
    assert "第三段也在" in texts


def test_bad_zip_fail_closed():
    check = verify_office_bytes(b"not-a-zip", ".docx")
    assert check["ok"] is False
    assert check["valid_ooxml"] is False
    with pytest.raises(ValueError, match="不是真 Word"):
        inspect_office_bytes(b"not-a-zip", ".docx")


def test_xlsx_kind_needs_sheet():
    with pytest.raises(ValueError, match="sheet"):
        parse_spec({"kind": "xlsx", "blocks": [{"type": "para", "text": "x"}]})


def test_new_tools_on_pi_allowlist():
    for name in ("render_document", "inspect_document", "verify_document"):
        assert name in ALLOWED_GATEWAY_TOOLS
    assert "bash" not in ALLOWED_GATEWAY_TOOLS


def test_plain_pptx_three_slides():
    body = "页一\n要点\n\n---\n页二\n要点\n\n---\n页三\n要点"
    raw = build_pptx_document(title="课.pptx", marker="mk-ppt", body=body)
    assert is_valid_ooxml_package(raw, ".pptx")
    outline = inspect_office_bytes(raw, ".pptx")
    assert outline["slides"] >= 3
